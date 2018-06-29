import argparse
import shutil, os
from html.parser import HTMLParser
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import subprocess

class MyHTMLParser(HTMLParser):

    def init(self, docfile):   
        self.data=""
        self.document=None
        self.flag=True
        self.total_len=0
        self.docfile = docfile
        self.document = Document()
        self.paragraph = self.document.add_paragraph()
        
    def handle_starttag(self, tag, attrs):
        #print("Encountered a start tag:", tag)
        self.data=""
        if tag == "p":
            self.data=""
            #print("Encountered a start tag:", tag)
        

    def handle_endtag(self, tag):
        #print("Encountered an end tag :", tag)
        #self.paragraph.add_run(self.data)
        if tag == "p" or tag == "br":
            #print(self.data)
            self.paragraph.add_run('\n')
            #paragraph = MyHTMLParser.document.add_paragraph(MyHTMLParser.data)
        

    def handle_data(self, data):
        #print("Encountered some data  :", data)
        #self.data += data
        self.paragraph.add_run(data)
        self.total_len += len(data)
        if self.flag:
            if self.total_len >= 300:
                self.flag = False
                self.write_middle_data()
        
    def set_files(self, docfile):
        self.docfile = docfile
        self.document = Document()
        self.paragraph = self.document.add_paragraph()

    def save_file(self):
        if self.flag:
            self.write_middle_data()
        self.document.save(self.docfile)

    def write_middle_data(self):
        pic = self.document.add_picture('image.jpg')
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.paragraph = self.document.add_paragraph()
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #self.document.add_picture('image.jpg')
        run = self.paragraph.add_run('Hire Me On\nwww.helpingtutors.com \nOr, Contact me\n')
        run.bold = True
        font = run.font
        font.size = Pt(16)
        run = self.paragraph.add_run('Hwhelp96@gmail.com \nI can send it’s answer instantly.\n')
        run.bold = True
        font = run.font
        font.size = Pt(16)
        self.paragraph = self.document.add_paragraph()
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
SRC=""

def list_files(directory, ext):
    files=os.listdir(directory)
    fx=[f for f in files if f.endswith('.' + ext)]
    return fx

def list_dirs(directory):
    dirs = os.listdir(directory)
    # fx=[ directory+'/'+f for f in dirs if os.path.isdir(directory+'/'+f)]
    fx = [f for f in dirs if os.path.isdir(directory + '/' + f)]
    return fx

def doc2pdf(docfile, pdffile):
    cmd = ['docto.exe', '-T', 'wdFormatPDF', '-f']
    cmd.append(docfile)
    cmd.append('-O')
    cmd.append(pdffile)
    returned_output = subprocess.check_output(cmd)
    
def main():
    #do for all directories in SRC
    dirs = list_dirs(SRC)

    for d in dirs:
        files = list_files(SRC+'/'+d, 'html')
        os.mkdir(d)
        pdf_dir = d + '/' + 'pdf'
        doc_dir = d + '/' + 'doc'
        os.makedirs(pdf_dir)
        os.makedirs(doc_dir)

        #do for all the html files in directory
        for f in files:
            src_file=SRC+'/'+d+'/'+f

            dst_pdf = pdf_dir + '/' + os.path.splitext(f)[0] + '.pdf'
            dst_doc = doc_dir + '/' + os.path.splitext(f)[0] + '.docx'

            print(".",end='',flush=True)
            #print(dst_pdf+' '+dst_doc)
            #move(src_file, dst)

            html = open(src_file).read()
            parser = MyHTMLParser()
            parser.init(dst_doc)
            parser.feed(html)
            parser.save_file()
            doc2pdf(dst_doc, dst_pdf)
            os.remove(src_file)
       

if __name__ == "__main__":
    """
    Execution starts here.
    """

    parser = argparse.ArgumentParser(description='Convert HTML files to PDF & Doc.')
    parser.add_argument('-s', '--src', help='Source directory', type=str, required=True)
    args = parser.parse_args()

    SRC=args.src
    main()

    '''
    a="""<p style="margin-left:-14.2pt;">The circuit diagram shown in Fig 3 is a differential amplifier, constructed using identical transistors. The device parameters are given as V<sub>BE</sub> = 0.7 V, V<sub>T</sub> = 0.025 V and ? = 120.</p> <p>a) Determine the biasing current, I<sub>C3</sub> and the d.c. output voltage at the output terminals V<sub>01</sub> and V<sub>02</sub> .</p> <p style="margin-left:-14.2pt;"></p> <p>b) Given the input signals: v<sub>i1</sub><strong>-</strong> v<sub>i2</sub> = 2sin(?t) (mV), determine the a.c. output voltage at the output terminal V<sub>01</sub> with respect to the ground.</p> <p>c) Given that the CMRR of the amplifier is 120 and the common-mode noise voltage at the input of the amplifier is 0.5mV, estimate the noise voltage at the output terminal V<sub>01</sub> and the signal-to-noise ratio at this terminal.</p> <p style="margin-left:-14.2pt;"></p> <p>d) Replacing the transistors in the differential amplifier of Fig 3 with Darlington pairs, sketch the circuit diagram of your modification and comment on the possible effect on the differential input impedances of the amplifier.</p>"""

    parser = MyHTMLParser()
    parser.init("15.docx")
    parser.feed(a)
    parser.save_file()
    doc2pdf("15.docx", "15.pdf")
    
    f=open("15.html").read()
    parser = MyHTMLParser()
    parser.feed(f)
    '''


    
